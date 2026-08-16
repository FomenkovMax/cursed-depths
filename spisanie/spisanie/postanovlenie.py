"""Чтение постановления об окончании ИП и извлечение из него реквизитов.

Поддерживаются .pdf (с текстовым слоем), .docx и .txt. Скан без текстового слоя
осознанно НЕ распознаётся: OCR даёт ошибки в цифрах, а здесь от цифр зависит
списание. Такой файл помечается как нечитаемый и уходит на ручное рассмотрение.
"""

from __future__ import annotations

import datetime as dt
import re
from dataclasses import dataclass, field
from pathlib import Path

from .common import norm_ip_nomer, norm_text, parse_date

DATE_RE = r"(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4})"

# Формулировки, подтверждающие отсутствие имущества и доходов.
# Список намеренно вынесен наверх: у разных ОСП формулировки отличаются,
# и дополнять его должен методолог, а не программист.
NO_PROPERTY_MARKERS = [
    "отсутствует имущество, на которое может быть обращено взыскание",
    "отсутствует имущество на которое может быть обращено взыскание",
    "меры по отысканию его имущества оказались безрезультатными",
    "меры по отысканию имущества оказались безрезультатными",
    "невозможно установить местонахождение должника, его имущества",
    "сведения о наличии принадлежащих ему денежных средств",
    "денежные средства и иное имущество отсутствуют",
    "должник не трудоустроен",
    "сведения о доходах должника отсутствуют",
    "источники дохода не установлены",
]


class UnreadableDocument(Exception):
    """Файл не прочитать: нет текстового слоя или формат не поддерживается."""


@dataclass
class Postanovlenie:
    """Реквизиты, извлечённые из постановления."""

    path: Path
    text: str
    ip_nomer: str | None = None
    data_vozbuzhdeniya: dt.date | None = None
    data_okonchaniya: dt.date | None = None
    fio: str | None = None
    data_rozhdeniya: dt.date | None = None
    punkt: int | None = None          # пункт ч. 1 ст. 46
    st46: bool = False                # в тексте есть ссылка на ст. 46
    net_imushchestva: bool = False    # найден маркер отсутствия имущества/доходов
    missing: list[str] = field(default_factory=list)  # что не удалось извлечь


# ---------------------------------------------------------------------------
# Чтение файла
# ---------------------------------------------------------------------------

def read_text(path: Path) -> str:
    suffix = path.suffix.lower()

    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    if suffix == ".pdf":
        import pymupdf

        with pymupdf.open(str(path)) as document:
            text = "\n".join(page.get_text() for page in document)
        # У скана текстовый слой пустой или состоит из мусора в пару символов.
        if len(text.strip()) < 80:
            raise UnreadableDocument(
                "PDF без текстового слоя (вероятно скан) — требуется распознавание"
            )
        return text

    raise UnreadableDocument(f"формат {suffix or '—'} не поддерживается")


# ---------------------------------------------------------------------------
# Извлечение реквизитов
# ---------------------------------------------------------------------------

def _search(patterns: list[str], text: str) -> str | None:
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(1).strip()
    return None


def parse(path: Path) -> Postanovlenie:
    """Прочитать постановление и вытащить реквизиты. Может бросить UnreadableDocument."""
    raw = read_text(path)
    text = re.sub(r"[ \t ]+", " ", raw)
    flat = norm_text(text)

    result = Postanovlenie(path=path, text=raw)

    result.ip_nomer = _search(
        [
            r"исполнительн\w*\s+производств\w*\s*(?:№|N|no)\s*([\d]+/[\d]+/[\d]+\s*-?\s*ИП)",
            r"(?:№|N)\s*([\d]+/[\d]+/[\d]+\s*-?\s*ИП)",
            r"([\d]{3,}/\d{2}/\d{3,}\s*-?\s*ИП)",
        ],
        text,
    )

    result.data_vozbuzhdeniya = parse_date(
        _search(
            [
                r"возбужден\w*\s*(?:исполнительное производство\s*)?(?:от\s*)?" + DATE_RE,
                r"дата возбуждения[^\d]{0,40}" + DATE_RE,
                r"производство[^.]{0,80}возбужден\w*[^\d]{0,20}" + DATE_RE,
            ],
            text,
        )
    )

    result.data_okonchaniya = parse_date(
        _search(
            [
                r"окончани\w*\s+исполнительного производства[^\d]{0,60}" + DATE_RE,
                r"окончить исполнительное производство[^\d]{0,60}" + DATE_RE,
                r"дата окончания[^\d]{0,40}" + DATE_RE,
                r"постановление[^\d]{0,40}от\s*" + DATE_RE,
                r"«?\s*(\d{1,2})\s*»?\s*\w+\s+\d{4}\s*г",
            ],
            text,
        )
    )

    result.fio = _search(
        [
            r"должник\w*\s*[:\-—]?\s*([А-ЯЁ][а-яё\-]+\s+[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)",
            r"в отношении\s+([А-ЯЁ][а-яё\-]+\s+[А-ЯЁ][а-яё\-]+(?:\s+[А-ЯЁ][а-яё\-]+)?)",
        ],
        text,
    )

    result.data_rozhdeniya = parse_date(
        _search(
            [
                r"дат\w*\s+рождения\s*[:\-—]?\s*" + DATE_RE,
                DATE_RE + r"\s*г\.?\s*р\.",
                r"\d{1,2}\.\d{1,2}\.\d{4}\s*года рождения",
            ],
            text,
        )
    )

    # \w* после «ст» — чтобы ловились все падежи: статья / статьи / статьёй / ст.
    result.st46 = bool(re.search(r"ст(?:ать\w*|\.)?\s*46", flat))
    punkt = _search([r"п(?:ункт\w*|\.)?\s*(\d)\s*ч(?:аст\w*|\.)?\s*1\s*ст(?:ать\w*|\.)?\s*46"], flat)
    if punkt:
        result.punkt = int(punkt)

    result.net_imushchestva = any(marker in flat for marker in NO_PROPERTY_MARKERS)

    for name, value in (
        ("номер ИП", result.ip_nomer),
        ("дата возбуждения ИП", result.data_vozbuzhdeniya),
        ("дата окончания ИП", result.data_okonchaniya),
        ("ФИО должника", result.fio),
        ("дата рождения должника", result.data_rozhdeniya),
    ):
        if not value:
            result.missing.append(name)

    return result


# ---------------------------------------------------------------------------
# Папка с постановлениями
# ---------------------------------------------------------------------------

SUPPORTED = {".pdf", ".docx", ".txt"}


def load_folder(folder: Path) -> tuple[dict[str, Postanovlenie], list[tuple[Path, str]]]:
    """Прочитать все постановления из папки.

    Возвращает (индекс по нормализованному номеру ИП, список нечитаемых файлов).
    """
    index: dict[str, Postanovlenie] = {}
    broken: list[tuple[Path, str]] = []

    for path in sorted(folder.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED:
            continue
        try:
            document = parse(path)
        except UnreadableDocument as exc:
            broken.append((path, str(exc)))
            continue
        except Exception as exc:  # noqa: BLE001 — один битый файл не должен ронять прогон
            broken.append((path, f"ошибка чтения: {exc}"))
            continue

        if not document.ip_nomer:
            broken.append((path, "в тексте не найден номер исполнительного производства"))
            continue

        key = norm_ip_nomer(document.ip_nomer)
        if key in index:
            broken.append((path, f"дубль постановления по ИП {document.ip_nomer}"))
            continue
        index[key] = document

    return index, broken
